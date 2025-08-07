from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from PIL import Image
import pytesseract
import speech_recognition as sr
import google.generativeai as genai
from docx import Document
from docx.shared import Inches
import os
import zipfile
import re
from datetime import datetime
import tempfile
import io
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import PyPDF2
from io import BytesIO
import secrets
import logging
import traceback

# Try to import docx2pdf, but make it optional
try:
    from docx2pdf import convert
    DOCX_TO_PDF_AVAILABLE = True
except ImportError:
    DOCX_TO_PDF_AVAILABLE = False
    print("Warning: docx2pdf not available. PDF conversion will be disabled.")

app = Flask(__name__)
CORS(app)  # Enable CORS for React frontend

# Setup logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# ----------- Configuration -----------
TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
OUTPUT_FOLDER = "generated_reports"
GEMINI_API_KEY = "AIzaSyBqVOV8quNcRTp03PgY7IjOmVsjHhDpEd8"

# SMTP Configuration
SMTP_CONFIG = {
    'host': 'smtp.gmail.com',
    'port': 587,
    'sender_email': 'aryavjain1708@gmail.com',
    'sender_password': 'whwkkfyvdzjzmwuu'
}

# ----------- Setup -----------
def setup_tesseract():
    """Setup Tesseract with error handling"""
    try:
        if os.path.exists(TESSERACT_PATH):
            pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
            logger.info("Tesseract configured successfully")
        else:
            logger.warning(f"Tesseract not found at {TESSERACT_PATH}")
            # Try default path
            pytesseract.pytesseract.tesseract_cmd = 'tesseract'
    except Exception as e:
        logger.error(f"Failed to setup Tesseract: {e}")

def setup_gemini():
    """Setup Gemini AI with error handling"""
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel('gemini-2.0-flash-exp')  # Using experimental model
        logger.info("Gemini AI configured successfully")
        return model
    except Exception as e:
        logger.error(f"Failed to setup Gemini AI: {e}")
        return None

setup_tesseract()
model = setup_gemini()

# ----------- Helper Functions -----------

def audio_to_text(audio_file):
    """Extract text from audio file"""
    try:
        logger.info(f"Processing audio file: {audio_file.filename}")
        
        # Save audio file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_audio:
            audio_file.save(temp_audio.name)
            temp_audio_path = temp_audio.name
        
        recognizer = sr.Recognizer()
        with sr.AudioFile(temp_audio_path) as source:
            audio = recognizer.record(source)
        
        text = recognizer.recognize_google(audio)
        logger.info("Audio processed successfully")
        
        # Clean up temp file
        os.unlink(temp_audio_path)
        
        return text
    except Exception as e:
        logger.error(f"Audio processing error: {e}")
        if 'temp_audio_path' in locals() and os.path.exists(temp_audio_path):
            os.unlink(temp_audio_path)
        return f"[Audio Error] {str(e)}"

def screenshot_to_text(image_file):
    """Extract text from image using OCR"""
    try:
        logger.info(f"Processing image file: {image_file.filename}")
        
        # Save image temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_image:
            image_file.save(temp_image.name)
            temp_image_path = temp_image.name
        
        image = Image.open(temp_image_path)
        text = pytesseract.image_to_string(image)
        logger.info("Image processed successfully")
        
        # Clean up temp file
        os.unlink(temp_image_path)
        
        return text
    except Exception as e:
        logger.error(f"Image processing error: {e}")
        if 'temp_image_path' in locals() and os.path.exists(temp_image_path):
            os.unlink(temp_image_path)
        return f"[Image Error] {str(e)}"

def generate_narrative_draft(audio_text="", screenshot_text="", extra_note=""):
    """Generate incident report using Gemini AI"""
    if not model:
        return "**Incident Report**\n[Error: Gemini AI not available]"
    
    prompt = f"""
Create a professional workplace harassment incident report with these exact section headers:
**Incident Report**
**Date:** {datetime.now().strftime('%Y-%m-%d')}
**Reporting Party:** Anonymous
**Description of Incident:**
[Detailed description based on provided evidence]
**Evidence Summary:**
[Summary of provided evidence]
**Requested Action:**
[Recommended actions]

Evidence provided:
1. Audio Transcript: {audio_text if audio_text else "None provided"}
2. Screenshot Text: {screenshot_text if screenshot_text else "None provided"}
3. Additional Notes: {extra_note if extra_note else "None provided"}

Please create a professional, detailed report based on this evidence.
"""
    try:
        logger.info("Generating narrative with Gemini AI")
        response = model.generate_content(prompt)
        logger.info("Narrative generated successfully")
        return response.text
    except Exception as e:
        logger.error(f"Narrative generation error: {e}")
        return f"**Incident Report**\n**Date:** {datetime.now().strftime('%Y-%m-%d')}\n[Error generating report: {str(e)}]"

def create_protected_pdf(docx_bytes, password):
    """Convert DOCX to password-protected PDF"""
    if not DOCX_TO_PDF_AVAILABLE:
        raise Exception("PDF conversion not available - docx2pdf not installed")
    
    try:
        # Save DOCX to temp file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_bytes)
            temp_docx_path = temp_docx.name
        
        # Create temp PDF path
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name
        
        # Convert DOCX to PDF
        convert(temp_docx_path, temp_pdf_path)
        
        # Password protect the PDF
        reader = PyPDF2.PdfReader(temp_pdf_path)
        writer = PyPDF2.PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page)
        
        writer.encrypt(password)
        
        protected_pdf = BytesIO()
        writer.write(protected_pdf)
        protected_pdf.seek(0)
        
        # Clean up temp files
        os.unlink(temp_docx_path)
        os.unlink(temp_pdf_path)
        
        return protected_pdf
        
    except Exception as e:
        # Clean up if error occurs
        if 'temp_docx_path' in locals() and os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)
        if 'temp_pdf_path' in locals() and os.path.exists(temp_pdf_path):
            os.unlink(temp_pdf_path)
        raise Exception(f"Failed to create protected PDF: {str(e)}")

def add_formatted_paragraph(doc, text):
    """Add paragraph to doc with **bold** formatting."""
    try:
        paragraph = doc.add_paragraph()
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)
    except Exception as e:
        logger.error(f"Error formatting paragraph: {e}")
        # Fallback: add as plain text
        doc.add_paragraph(text)

def create_report_document(narrative_text, audio_filename=None, image_filename=None, image_file=None):
    """Create DOCX document with the report"""
    try:
        doc = Document()
        doc.add_heading('Workplace Harassment Incident Report', level=1)

        doc.add_heading('Generated Narrative', level=2)
        for line in narrative_text.strip().split("\n"):
            if line.strip():
                add_formatted_paragraph(doc, line.strip())

        # Handle image attachment
        if image_file:
            doc.add_heading('Attached Screenshot', level=2)
            try:
                # Save image temporarily for embedding
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_img:
                    image_file.seek(0)  # Reset file pointer
                    temp_img.write(image_file.read())
                    temp_img_path = temp_img.name
                
                doc.add_picture(temp_img_path, width=Inches(5.5))
                os.unlink(temp_img_path)  # Clean up
            except Exception as e:
                logger.error(f"Failed to embed image: {e}")
                doc.add_paragraph(f"[Screenshot could not be embedded: {e}]")
        elif image_filename:
            doc.add_paragraph(f"[Screenshot attached separately: {image_filename}]")

        if audio_filename:
            doc.add_heading('Attached Audio Evidence', level=2)
            doc.add_paragraph(f"[Audio File Attached Separately: {audio_filename}]")

        logger.info("Document created successfully")
        return doc
    except Exception as e:
        logger.error(f"Error creating document: {e}")
        raise

def send_email_with_attachment(subject, body, to_email, attachment_bytes=None, filename=None, password=None):
    """Send email with optional attachment"""
    try:
        msg = MIMEMultipart()
        msg["From"] = SMTP_CONFIG['sender_email']
        msg["To"] = to_email
        msg["Subject"] = subject

        msg.attach(MIMEText(body, "plain"))

        if attachment_bytes and filename:
            part = MIMEApplication(attachment_bytes, Name=filename)
            part['Content-Disposition'] = f'attachment; filename="{filename}"'
            msg.attach(part)

            if password:
                msg.attach(MIMEText(f"\n\nProtected PDF Password: {password}", "plain"))

        server = smtplib.SMTP(SMTP_CONFIG['host'], SMTP_CONFIG['port'])
        server.starttls()
        server.login(SMTP_CONFIG['sender_email'], SMTP_CONFIG['sender_password'])
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        logger.error(f"Failed to send email: {str(e)}")
        return False

def assess_severity(narrative_text):
    """Assess the severity of the incident using Gemini AI"""
    if not model:
        return {
            'severity': 'MODERATE', 
            'rationale': 'AI model not available - defaulting to moderate severity for manual review.',
            'actions_taken': ['Report logged for manual review', 'Notification sent to HR team']
        }
    
    prompt = f"""
Analyze this workplace harassment incident report and provide a comprehensive assessment.

Classification Criteria:
- HIGH: Physical violence, sexual harassment, threats, discriminatory actions
- MODERATE: Verbal harassment, hostile work environment, inappropriate behavior
- LOW: Minor workplace conflicts, misunderstandings, first-time minor incidents

Report to analyze:
{narrative_text}

Please provide your assessment in this exact format:

SEVERITY: [HIGH|MODERATE|LOW]

RATIONALE: [Provide a detailed 2-3 sentence explanation of why this incident falls into this severity category. Be specific about which behaviors/actions led to this classification.]

RECOMMENDED_ACTIONS: [List 2-3 specific recommended actions based on the severity level]

SUPPORT_RESOURCES: [Suggest 1-2 support resources or next steps for the victim]
"""
    
    try:
        logger.info("Generating comprehensive severity assessment")
        response = model.generate_content(prompt)
        result = response.text
        
        # Parse the response
        severity = "MODERATE"  # Default
        rationale = "Unable to determine severity from the provided information."
        actions_taken = []
        support_resources = []
        
        # Extract severity
        if "SEVERITY: HIGH" in result:
            severity = "HIGH"
        elif "SEVERITY: LOW" in result:
            severity = "LOW"
        elif "SEVERITY: MODERATE" in result:
            severity = "MODERATE"
        
        # Extract rationale
        if "RATIONALE:" in result:
            rationale_part = result.split("RATIONALE:")[1]
            if "RECOMMENDED_ACTIONS:" in rationale_part:
                rationale = rationale_part.split("RECOMMENDED_ACTIONS:")[0].strip()
            elif "SUPPORT_RESOURCES:" in rationale_part:
                rationale = rationale_part.split("SUPPORT_RESOURCES:")[0].strip()
            else:
                rationale = rationale_part.strip()
        
        # Extract recommended actions
        if "RECOMMENDED_ACTIONS:" in result:
            actions_part = result.split("RECOMMENDED_ACTIONS:")[1]
            if "SUPPORT_RESOURCES:" in actions_part:
                actions_text = actions_part.split("SUPPORT_RESOURCES:")[0].strip()
            else:
                actions_text = actions_part.strip()
            
            # Parse actions (assuming they're in a list format or separated by newlines)
            actions_taken = [action.strip().lstrip('- ').lstrip('* ') for action in actions_text.split('\n') if action.strip()]
            actions_taken = [action for action in actions_taken if action and len(action) > 5]  # Filter out empty/short lines
        
        # Extract support resources
        if "SUPPORT_RESOURCES:" in result:
            support_text = result.split("SUPPORT_RESOURCES:")[1].strip()
            support_resources = [resource.strip().lstrip('- ').lstrip('* ') for resource in support_text.split('\n') if resource.strip()]
            support_resources = [resource for resource in support_resources if resource and len(resource) > 5]
        
        # Add default actions if none were extracted
        if not actions_taken:
            if severity == "HIGH":
                actions_taken = [
                    "Immediate escalation to senior management",
                    "HR investigation initiated within 24 hours",
                    "Victim safety measures implemented"
                ]
            elif severity == "MODERATE":
                actions_taken = [
                    "Formal HR investigation scheduled",
                    "Documentation preserved for review",
                    "Mediation or disciplinary action considered"
                ]
            else:  # LOW
                actions_taken = [
                    "Incident logged for monitoring",
                    "Informal resolution attempted",
                    "Follow-up scheduled if behavior continues"
                ]
        
        # Add default support resources if none were extracted
        if not support_resources:
            support_resources = [
                "Employee Assistance Program (EAP) counseling available",
                "HR open-door policy for ongoing support"
            ]
        
        logger.info(f"Assessment completed - Severity: {severity}")
        return {
            'severity': severity, 
            'rationale': rationale,
            'actions_taken': actions_taken,
            'support_resources': support_resources
        }
        
    except Exception as e:
        logger.error(f"Severity assessment error: {e}")
        return {
            'severity': 'MODERATE', 
            'rationale': f'Assessment error occurred: {str(e)}. Defaulting to moderate severity for manual review.',
            'actions_taken': ['Report flagged for immediate manual review', 'Technical support notified'],
            'support_resources': ['Contact HR directly for immediate assistance']
        }

# ----------- Routes -----------

@app.route('/api/generate-report', methods=['POST'])
def generate_report():
    """Generate incident report endpoint"""
    try:
        logger.info("Starting report generation")
        logger.info(f"Request files: {list(request.files.keys())}")
        logger.info(f"Request form: {dict(request.form)}")
        
        audio_file = request.files.get('audio')
        image_file = request.files.get('image')
        additional_notes = request.form.get('additionalNotes', '')
        
        logger.info(f"Audio file: {audio_file.filename if audio_file else 'None'}")
        logger.info(f"Image file: {image_file.filename if image_file else 'None'}")
        logger.info(f"Additional notes: {additional_notes}")
        
        # Process audio and image
        audio_text = ""
        screenshot_text = ""
        
        if audio_file:
            audio_text = audio_to_text(audio_file)
            logger.info(f"Audio text extracted: {len(audio_text)} characters")
        
        if image_file:
            screenshot_text = screenshot_to_text(image_file)
            logger.info(f"Screenshot text extracted: {len(screenshot_text)} characters")
        
        # Generate narrative
        logger.info("Generating narrative")
        narrative = generate_narrative_draft(audio_text, screenshot_text, additional_notes)
        
        # Create DOCX
        logger.info("Creating document")
        doc = create_report_document(
            narrative,
            audio_file.filename if audio_file else None,
            image_file.filename if image_file else None,
            image_file
        )
        
        # Save DOCX to memory
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        # Create ZIP with attachments
        logger.info("Creating ZIP file")
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zipf:
            zipf.writestr('incident_report.docx', doc_buffer.getvalue())
            if audio_file:
                audio_file.seek(0)  # Reset file pointer
                zipf.writestr(f'audio_evidence.{audio_file.filename.split(".")[-1]}', audio_file.read())
            if image_file:
                image_file.seek(0)  # Reset file pointer
                zipf.writestr(f'image_evidence.{image_file.filename.split(".")[-1]}', image_file.read())
        
        zip_buffer.seek(0)
        
        logger.info("Report generation completed successfully")
        return jsonify({
            'success': True,
            'narrative_preview': narrative[:500] + ("..." if len(narrative) > 500 else ""),
            'audio_text_preview': audio_text[:200] + ("..." if len(audio_text) > 200 else "") if audio_text and not audio_text.startswith('[Audio Error]') else None,
            'screenshot_text_preview': screenshot_text[:200] + ("..." if len(screenshot_text) > 200 else "") if screenshot_text and not screenshot_text.startswith('[Image Error]') else None,
            'zip_data': zip_buffer.getvalue().hex()
        })
        
    except Exception as e:
        logger.error(f"Error in generate_report: {e}")
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/assess-severity', methods=['POST'])
def assess_severity_route():
    """Assess severity and send emails with attachments"""
    try:
        data = request.json
        narrative_text = data.get('narrative_text')
        zip_data_hex = data.get('zip_data')
        
        if not narrative_text or not zip_data_hex:
            return jsonify({'success': False, 'error': 'Missing data'}), 400
        
        # Get all files from ZIP data
        zip_data = BytesIO(bytes.fromhex(zip_data_hex))
        with zipfile.ZipFile(zip_data) as zipf:
            # Get DOCX
            with zipf.open('incident_report.docx') as docx_file:
                docx_bytes = docx_file.read()
            
            # Get audio file if exists
            audio_bytes = None
            audio_filename = None
            for filename in zipf.namelist():
                if filename.startswith('audio_evidence.'):
                    with zipf.open(filename) as audio_file:
                        audio_bytes = audio_file.read()
                        audio_filename = filename
                    break
        
        # Assess severity
        assessment = assess_severity(narrative_text)
        
        # Create protected PDF if available
        if DOCX_TO_PDF_AVAILABLE:
            password = secrets.token_urlsafe(12)
            protected_pdf = create_protected_pdf(docx_bytes, password)
            pdf_bytes = protected_pdf.getvalue()
            pdf_filename = "incident_report.pdf"
        else:
            password = None
            pdf_bytes = docx_bytes
            pdf_filename = "incident_report.docx"
        
        # Determine recipient
        if assessment['severity'] == "HIGH":
            recipient = "bluecndy18@gmail.com"
            subject = "🚨 Urgent: High Severity Incident Report"
        else:
            recipient = "meetkadam420@gmail.com"
            subject = "⚠️ Incident Report Notification"
        
        # Create email with attachments
        msg = MIMEMultipart()
        msg["From"] = SMTP_CONFIG['sender_email']
        msg["To"] = recipient
        msg["Subject"] = subject

        # Email body
        email_body = f"""WORKPLACE HARASSMENT INCIDENT REPORT

SEVERITY ASSESSMENT: {assessment['severity']}

DETAILED RATIONALE:
{assessment['rationale']}

RECOMMENDED ACTIONS:
{chr(10).join(f"• {action}" for action in assessment.get('actions_taken', []))}

SUPPORT RESOURCES:
{chr(10).join(f"• {resource}" for resource in assessment.get('support_resources', []))}

{f'PDF Password: {password}' if password else ''}
"""
        msg.attach(MIMEText(email_body, "plain"))

        # Attach PDF
        part = MIMEApplication(pdf_bytes, Name=pdf_filename)
        part['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
        msg.attach(part)

        # Attach audio if exists
        if audio_bytes and audio_filename:
            audio_part = MIMEApplication(audio_bytes, Name=audio_filename)
            audio_part['Content-Disposition'] = f'attachment; filename="{audio_filename}"'
            msg.attach(audio_part)

        # Send email
        try:
            server = smtplib.SMTP(SMTP_CONFIG['host'], SMTP_CONFIG['port'])
            server.starttls()
            server.login(SMTP_CONFIG['sender_email'], SMTP_CONFIG['sender_password'])
            server.send_message(msg)
            server.quit()
            email_sent = True
        except Exception as e:
            logger.error(f"Email sending failed: {e}")
            email_sent = False
        
        return jsonify({
            'success': True,
            'severity': assessment['severity'],
            'email_sent': email_sent,
            'attachments': {
                'pdf': True,
                'audio': audio_bytes is not None
            },
            'recipient': recipient
        })
        
    except Exception as e:
        logger.error(f"Error in assess_severity: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/download-report', methods=['POST'])
def download_report():
    """Download report endpoint"""
    try:
        zip_data_hex = request.json.get('zip_data')
        if not zip_data_hex:
            return jsonify({'error': 'No data provided'}), 400
        
        zip_data = BytesIO(bytes.fromhex(zip_data_hex))
        return send_file(
            zip_data,
            mimetype='application/zip',
            as_attachment=True,
            download_name='incident_report.zip'
        )
    except Exception as e:
        logger.error(f"Error in download_report: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({
        'status': 'healthy',
        'tesseract_available': os.path.exists(TESSERACT_PATH),
        'gemini_available': model is not None,
        'docx2pdf_available': DOCX_TO_PDF_AVAILABLE
    })

if __name__ == '__main__':
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    logger.info("Starting Flask application")
    app.run(debug=True, port=5005)